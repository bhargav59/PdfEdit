"""Generate test PDF and DOCX fixtures for the test suite."""

import os
import sys

# Ensure reportlab and python-docx are available
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except ImportError:
    print("reportlab not installed. Run: pip install reportlab")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))


def create_pdf(filename: str, num_pages: int) -> None:
    path = os.path.join(FIXTURES_DIR, filename)
    c = canvas.Canvas(path, pagesize=A4)
    for i in range(1, num_pages + 1):
        c.setFont("Helvetica", 24)
        c.drawString(100, 700, f"Test Page {i}")
        c.setFont("Helvetica", 12)
        c.drawString(100, 670, f"This is page {i} of {num_pages}")
        if i < num_pages:
            c.showPage()
    c.save()
    print(f"Created {path} ({num_pages} pages)")


def create_docx(filename: str) -> None:
    path = os.path.join(FIXTURES_DIR, filename)
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test Word document for conversion testing.")
    doc.add_paragraph("It contains multiple paragraphs of sample text.")
    doc.save(path)
    print(f"Created {path}")


def create_corrupt_pdf(filename: str) -> None:
    path = os.path.join(FIXTURES_DIR, filename)
    with open(path, "wb") as f:
        f.write(b"%PDF-1.4 CORRUPT DATA " + os.urandom(200))
    print(f"Created {path} (corrupt)")


if __name__ == "__main__":
    create_pdf("1-page.pdf", 1)
    create_pdf("5-page.pdf", 5)
    create_docx("sample.docx")
    create_corrupt_pdf("corrupt.pdf")
    print("All fixtures generated.")
